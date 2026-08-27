from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "output/docs/CS2103T_Week_2_Lecture_Notes.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "E8EEF5")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(" " + body)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_document():
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CS2103/T Week 2 Lecture Notes")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Topics: SDLC process models, Git/GitHub, IDEs, and automated testing").italic = True
    add_para(doc, "Source: CS2103/T Week 2 Topics PDF, exported 21 Aug 2026. These notes reorganize and explain the source material in lecturer style for study.")

    add_heading(doc, "How to Study Week 2", 1)
    add_callout(
        doc,
        "Big picture.",
        "Week 2 moves from simply writing code to managing software work. The core question is: how do we control change when the product, team, and uncertainty grow?"
    )
    add_bullets(doc, [
        "First, understand process models: they explain how a team moves through requirements, design, coding, testing, release, and later maintenance.",
        "Second, treat Git and GitHub as a safety net and history reader, not just as commands to memorize.",
        "Third, use an IDE to reduce mechanical coding effort and to observe program behavior.",
        "Fourth, start testing early because every new feature can accidentally break existing behavior.",
    ])

    add_heading(doc, "W2.1 SDLC Process Models: Basics", 1)
    add_heading(doc, "Why Process Models Exist", 2)
    add_para(doc, "A very small software project can sometimes survive with code-and-fix: start coding, run the program, fix whatever breaks, and repeat. This has almost no planning overhead, so it can work for a short one-person program.")
    add_para(doc, "But as the program or team grows, code-and-fix becomes weak. The team cannot easily tell how far along the project is, divide work cleanly, or remember why earlier choices were made. Changes become more expensive because there is no shared roadmap.")
    add_para(doc, "The software development lifecycle, or SDLC, names the main activities software usually passes through: requirements, analysis, design, implementation, testing, deployment, operation, and maintenance. Process models describe different ways to organize those activities.")

    add_callout(
        doc,
        "Lecturer explanation.",
        "Think of SDLC as the list of jobs that must happen; a process model is the route you choose through those jobs."
    )

    add_heading(doc, "Common SDLC Activities", 2)
    add_table(doc, ["Activity", "Purpose", "Typical output"], [
        ("Requirements", "Find out what the system should do and what constraints it must satisfy.", "Requirement statements, user needs, acceptance criteria."),
        ("Analysis", "Understand the problem, domain, and feasibility more deeply.", "Problem model, assumptions, risks, clarified scope."),
        ("Design", "Decide how the system will be structured before implementation.", "Architecture, component design, UI/data/API design."),
        ("Implementation", "Write and integrate the code.", "Working code, builds, commits."),
        ("Testing", "Run the system under specified conditions and compare with expected behavior.", "Test results, bug reports, confidence evidence."),
        ("Deployment", "Make the software available to users.", "Release package, production setup, delivery process."),
        ("Operation and maintenance", "Keep the software useful after release as users, defects, and environments change.", "Bug fixes, improvements, future work items."),
    ], [1.55, 2.8, 2.15])

    add_heading(doc, "Sequential / Waterfall Model", 2)
    add_para(doc, "The sequential model, often called the waterfall model, treats development as a linear movement through stages. A stage is completed, produces artifacts, and those artifacts become input for the next stage.")
    add_bullets(doc, [
        "Example: requirements are gathered first, then used by the design stage.",
        "In a strict sequential model, the project moves forward only; once a stage is finished, the team does not revise it later.",
        "In practice, teams often relax this and allow a later stage to send work back to the previous stage, but that means redoing work considered finished.",
    ])
    add_table(doc, ["Strength", "Why it helps"], [
        ("Easy stage-level tracking", "You can say the project is in requirements, design, implementation, or testing."),
        ("Clear artifacts", "Each stage has a visible output for the next stage."),
        ("Suitable for stable problems", "If requirements are well understood and unlikely to change, estimates can be more reliable."),
    ], [2.0, 4.5])
    add_table(doc, ["Weakness", "Why it hurts"], [
        ("Poor fit for unclear problems", "Users may not know their true needs until they see or use a working system."),
        ("Late feedback", "Integration and user contact may happen near the end, when changes cost more."),
        ("Hidden stage overrun", "Progress within a long stage can still be hard to see until the delay has already grown."),
    ], [2.0, 4.5])

    add_heading(doc, "Iterative Model", 2)
    add_para(doc, "The iterative model builds the software through several cycles. Each iteration can pass through requirements, design, implementation, testing, and even deployment. It produces a new product version or usable improvement.")
    add_para(doc, "Feedback from one iteration is used to improve later iterations. If an implementation task took longer than expected, later estimates can be adjusted. If a feature is not well received by users, it can be changed or removed.")
    add_callout(
        doc,
        "Key distinction.",
        "A sequential project is divided by activity. An iterative project is divided by bounded cycles, each designed to produce evidence and learning."
    )

    add_heading(doc, "Breadth-First and Depth-First Iterations", 2)
    add_table(doc, ["Approach", "Meaning", "Minesweeper example"], [
        ("Breadth-first", "Each iteration evolves most major components and feature areas in parallel, producing a working product each time.", "An early version is playable but primitive: text UI, fixed board size, and limited mine layouts."),
        ("Depth-first", "Each iteration focuses deeply on one component, feature area, risk, or assumption; early iterations may not produce a complete working product.", "One iteration builds the full UI without game logic; another focuses only on minefield generation."),
        ("Mixed", "A project can combine both styles depending on the risk and learning goal of each iteration.", "One iteration may improve the playable game while also exploring a risky algorithm."),
    ], [1.35, 2.65, 2.5])
    add_para(doc, "Iterating and incrementing are related but different. To iterate is to rework existing material using feedback. To increment is to add something new. Most real projects are both iterative and incremental.")
    add_para(doc, "The result of an iteration is an increment: a usable improvement or addition to the product, not merely a new code version.")
    add_callout(
        doc,
        "Evidence rule.",
        "Before an iteration starts, decide what success evidence should exist at the end: a passing test, a working demo, a target-user response, or a decision made from the result."
    )

    add_heading(doc, "Risk-Driven Iterations and the Spiral Idea", 2)
    add_para(doc, "An early iteration is valuable because it lets the team discover wrong assumptions while changing direction is still cheap. Therefore, risky assumptions should be tested early: unfamiliar technology, unclear user need, or a difficult performance target.")
    add_para(doc, "Ordering iterations by risk is the central idea of the spiral model. The team spirals through planning, building, and learning while reducing the biggest risks first.")

    add_heading(doc, "AI Impact on Process Models", 2)
    add_para(doc, "AI coding tools make candidate implementations cheaper and faster. However, deciding what to build and verifying that the result is correct remain difficult.")
    add_para(doc, "A vague requirement that might once have produced a clarifying question from a teammate can now produce a fast but wrong implementation. This makes precise specification and verification even more important.")

    add_heading(doc, "Sequential vs Iterative: Lecturer Comparison", 2)
    add_table(doc, ["Question", "Sequential answer", "Iterative answer"], [
        ("How is work divided?", "By activity: requirements, design, implementation, testing.", "By cycles: each iteration produces learning or an increment."),
        ("When does feedback arrive?", "Often late, especially user feedback.", "Earlier and repeatedly."),
        ("Best fit?", "Stable, well-understood requirements.", "Uncertain requirements, changing needs, risky assumptions."),
        ("Main danger?", "Discovering mistakes when revising is expensive.", "Running iterations without clear evidence or decisions."),
    ], [1.65, 2.4, 2.45])

    add_heading(doc, "W2.2 and W2.3 RCS: Git, GitHub, and Revision History", 1)
    add_heading(doc, "AI's Impact on Git and GitHub", 2)
    add_para(doc, "The Week 2 notes frame Git as something developers increasingly lean on rather than manually type from memory. AI can run commands for you, but you must still understand what those commands do.")
    add_bullets(doc, [
        "Remembering exact commands matters less than understanding their effect.",
        "Being able to go back matters more because AI can change many files quickly.",
        "Committing before risky work gives you a safe return point.",
        "Reading diffs and pull requests matters more because you may review code you did not personally write.",
        "Small commits with clear messages are easier to review, understand, and undo.",
    ])
    add_heading(doc, "Week 2 Git-Mastery Tours", 2)
    add_table(doc, ["Item", "Tour focus", "What you should be able to explain"], [
        ("W2.2a", "Tour 2: Backing up a repo on the cloud.", "Why a remote copy protects work and enables sharing."),
        ("W2.2b", "Tour 3: Working off a remote repo.", "How local and remote repositories relate during fetch/pull/push style workflows."),
        ("W2.3a", "Tour 4: Using revision history.", "How history helps inspect old states, understand changes, and trace decisions."),
        ("W2.3b", "Tour 5: Fine-tuning revision history.", "Why history can be tidied or shaped so it stays readable and useful."),
    ], [1.0, 2.25, 3.25])

    add_heading(doc, "W2.4 IDEs: Basic Features", 1)
    add_para(doc, "An Integrated Development Environment, or IDE, supports many development activities in one tool. It is not just a text editor; it brings editing, building, running, debugging, testing, and other support together.")
    add_table(doc, ["IDE part", "What it does", "Why it matters for beginners"], [
        ("Source code editor", "Syntax coloring, auto-completion, code navigation, error highlighting, and snippets.", "Helps you notice mistakes earlier and move around code faster."),
        ("Compiler/interpreter and build support", "Compiles, links, runs, and packages programs.", "Reduces friction when repeatedly running your project."),
        ("Debugger", "Runs the program step by step so you can inspect runtime behavior.", "Lets you see what the code actually does, not what you hoped it would do."),
        ("Additional tools", "Testing support, UI builders, version-control support, runtime simulation, modeling, AI assistance, and collaboration.", "Turns the IDE into a broader engineering workspace."),
    ], [1.6, 2.65, 2.25])
    add_heading(doc, "Examples of IDEs", 2)
    add_bullets(doc, [
        "Java: Eclipse, IntelliJ IDEA, NetBeans.",
        "C#, C++: Visual Studio.",
        "Swift: Xcode.",
        "Python: PyCharm.",
        "Multiple languages: VS Code.",
        "Some experienced developers, especially those with UNIX backgrounds, prefer lightweight but powerful editors such as Vim or NeoVim.",
    ])
    add_para(doc, "The source also points students to setup guides for IntelliJ IDEA and VS Code. In this course, the practical goal is to know enough IDE usage to work productively on the individual project.")

    add_heading(doc, "W2.5 Introduction to Automated Testing", 1)
    add_heading(doc, "What Testing Means", 2)
    add_para(doc, "The source uses the IEEE idea of testing: operate a system or component under specified conditions, observe or record the results, and evaluate some aspect of the system.")
    add_para(doc, "A test case specifies how to perform a test. At minimum, it gives the input to the software under test, or SUT, and the expected behavior.")
    add_callout(
        doc,
        "Lecturer explanation.",
        "A test case is a promise written before or during testing: if we do this input, the system should behave like that. Testing is the act of checking that promise."
    )
    add_heading(doc, "Test Case Example: Browser", 2)
    add_table(doc, ["Part", "Details"], [
        ("Input", "Start the browser using a blank page with the vertical scrollbar disabled. Then load longfile.html from the test data folder."),
        ("Expected behavior", "The scrollbar should be automatically enabled after longfile.html loads."),
        ("Failure", "If the scrollbar remains disabled, the test case fails."),
        ("Possible defect", "The underlying bug might be something like an uninitialized variable, although the test case itself could also be wrong."),
    ], [1.45, 5.05])
    add_heading(doc, "How to Execute a Test Case", 2)
    add_numbers(doc, [
        "Feed the specified input to the SUT.",
        "Observe the actual output or behavior.",
        "Compare the actual output with the expected output.",
    ])
    add_para(doc, "A test case failure is a mismatch between expected behavior and actual behavior. It indicates a potential defect. The word potential matters because the test itself may contain an incorrect expectation.")

    doc.add_page_break()
    add_heading(doc, "Regression Testing", 2)
    add_para(doc, "A regression is an unintended and undesirable effect caused by modifying a system. Regression testing means re-testing the software to detect such regressions.")
    add_para(doc, "The usual method is to retest all related components, even those already tested before. Regression testing is more effective when done frequently after small changes. Manual regression testing can become too expensive, so automation makes it practical.")
    add_table(doc, ["Concept", "Meaning", "Practical lesson"], [
        ("Regression", "A new change breaks existing behavior.", "Do not assume old features still work after adding a new feature."),
        ("Regression testing", "Retesting to detect regressions.", "Run tests repeatedly, especially after small changes."),
        ("Automation", "Tests are executed and judged programmatically.", "Makes frequent testing affordable and precise."),
    ], [1.5, 2.5, 2.5])

    add_heading(doc, "Automated Testing", 2)
    add_para(doc, "An automated test case can be run programmatically, and the pass/fail result is also determined programmatically. Compared with manual testing, automation reduces repeated effort and improves precision because manual testing is prone to human error.")
    add_para(doc, "The optional Week 2 item mentions automated testing of CLI applications. For this course project, that connects naturally to running command-line inputs and checking exact output.")
    add_heading(doc, "AI's Impact on Testing", 2)
    add_bullets(doc, [
        "AI reduces the effort needed to draft test cases, write test code, and set up testing mechanisms.",
        "Because automation is easier, not having time to write tests is a weaker excuse.",
        "AI cannot know expected behavior unless a human or specification defines it.",
        "If the same AI writes both implementation and tests, it may repeat the same misunderstanding in both.",
        "Automated regression tests become more valuable because AI agents can make large changes quickly.",
    ])

    doc.add_page_break()
    add_heading(doc, "Connections to Your iP Work", 1)
    add_table(doc, ["Week 2 idea", "How it appears in your project"], [
        ("Iterative development", "You completed Level 0 through Level 6 as small increments rather than building everything at once."),
        ("Git history", "Each level and extension can be tagged, committed, reviewed, and pushed."),
        ("IDE support", "IntelliJ or another IDE helps navigate classes such as Baby, Task, Todo, Deadline, Event, and Command."),
        ("Regression testing", "Your UI test plan reruns old behavior after each new feature such as mark, unmark, deadline, event, delete, and enum refactoring."),
        ("Automation", "The test-ui script runs command-line scenarios and compares exact expected output."),
    ], [1.65, 4.85])

    add_heading(doc, "Exam and Tutorial Style Questions", 1)
    add_bullets(doc, [
        "Explain why code-and-fix may work for a tiny one-person program but not for a larger team project.",
        "Compare sequential and iterative process models using feedback timing, suitability, and risk.",
        "Use Minesweeper to explain breadth-first and depth-first iterations.",
        "Define an increment and explain why a code version alone may not be a useful increment.",
        "Explain why AI coding makes specification and verification more important, not less important.",
        "Explain why small commits with clear messages are useful when using AI-generated changes.",
        "Name the main parts of an IDE and explain how each helps development.",
        "Define test case, SUT, failure, regression, regression testing, and automated test case.",
        "Explain why a failing test indicates a potential defect rather than guaranteed implementation bug.",
        "Explain why automated regression testing is especially valuable after small frequent changes.",
    ])

    add_heading(doc, "One-Page Recap", 1)
    add_bullets(doc, [
        "SDLC activities include requirements, analysis, design, implementation, testing, deployment, operation, and maintenance.",
        "Sequential models are linear and easy to track at stage level, but weak when requirements are uncertain or feedback arrives late.",
        "Iterative models build through cycles, use feedback, and can be breadth-first, depth-first, or mixed.",
        "Good iterations end in evidence that supports a decision.",
        "Risky assumptions should be tested early; this is the spirit of the spiral model.",
        "AI makes coding faster but increases the need for precise requirements, careful verification, readable Git history, and automated regression tests.",
        "IDEs integrate editing, building, running, debugging, testing, and other development support.",
        "Testing compares actual behavior with expected behavior; regression testing checks that changes did not break existing behavior.",
        "Automated tests reduce repeated effort and human error, but expected behavior still has to come from a trustworthy specification or human understanding.",
    ])

    add_heading(doc, "Source Coverage Map", 1)
    for item in [
        "W2.1a: SDLC introduction, code-and-fix, lifecycle activities, maintenance feedback.",
        "W2.1b: Sequential/waterfall model, artifacts, strengths, and weaknesses.",
        "W2.1c: Iterative model, increments, breadth-first/depth-first, risk, spiral model, AI impact.",
        "W2.2a-b: GitHub remote backup and working from a remote repo as Git-Mastery tours.",
        "W2.3a-b: Using and fine-tuning revision history as Git-Mastery tours.",
        "W2.4a-b: IDE definition, parts, examples, and setup references.",
        "W2.5a-d: Testing definition, test cases, regression testing, automation, CLI testing item, AI impact.",
    ]:
        label, detail = item.split(": ", 1)
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(label + ": ")
        run.bold = True
        p.add_run(detail)

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CS2103/T Week 2 Lecture Notes")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
